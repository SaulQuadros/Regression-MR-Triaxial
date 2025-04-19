#!/usr/bin/env python
# coding: utf-8

# --- app_latex.py ---
import io
from io import BytesIO
from docx import Document
from docx.shared import Inches

def add_formatted_equation(doc, eq_text):
    """
    Adiciona a equação ao Word, formatando:
    - '^' para sobrescrito
    - '_' ou '~' para subscrito
    - 'σ' seguido de caractere para subscrito
    """
    eq = eq_text.strip().strip("$$")
    p = doc.add_paragraph()
    i = 0
    while i < len(eq):
        ch = eq[i]
        if ch == '^':
            # superscrito
            i += 1
            exp = ""
            while i < len(eq) and (eq[i].isdigit() or eq[i] in ['.', '-']):
                exp += eq[i]
                i += 1
            run = p.add_run(exp)
            run.font.superscript = True
        elif ch in ['_', '~']:
            # subscrito
            i += 1
            if i < len(eq):
                run = p.add_run(eq[i])
                run.font.subscript = True
                i += 1
        elif ch == 'σ':
            # sigma + subscrito opcional
            p.add_run('σ')
            i += 1
            if i < len(eq) and (eq[i].isdigit() or eq[i].isalpha()):
                sub = eq[i]
                run_sub = p.add_run(sub)
                run_sub.font.subscript = True
                i += 1
        else:
            p.add_run(ch)
            i += 1
    return p

def add_data_table(doc, df):
    """
    Adiciona ao Word uma tabela com os dados do DataFrame df.
    """
    doc.add_heading("Dados do Ensaio Triaxial", level=2)
    table = doc.add_table(rows=df.shape[0] + 1, cols=df.shape[1])
    table.style = 'Light List Accent 1'
    # cabeçalho
    for j, col in enumerate(df.columns):
        table.rows[0].cells[j].text = str(col)
    # dados
    for i in range(df.shape[0]):
        for j, col in enumerate(df.columns):
            table.rows[i + 1].cells[j].text = str(df.iloc[i, j])
    return doc

def generate_word_doc(eq_latex, metrics_txt, fig, energy, degree, intercept, df):
    """
    Gera um documento .docx com a equação, métricas, tabela e gráfico 3D.
    """
    doc = Document()
    doc.add_heading("Relatório de Regressão", level=1)
    doc.add_heading("Configurações", level=2)
    doc.add_paragraph(f"Tipo de energia: {energy}")
    if degree is not None:
        doc.add_paragraph(f"Grau polinomial: {degree}")
        doc.add_heading("Equação Ajustada", level=2)

    # Equação formatada
    raw_eq = eq_latex.strip("$$")
    for ln in raw_eq.split("\\\\"):
        ln = ln.replace("σ₃", "σ_3").replace("σd", "σ_d")
        add_formatted_equation(doc, ln)

    # Indicadores
    doc.add_heading("Indicadores Estatísticos", level=2)
    doc.add_paragraph(metrics_txt)
    doc.add_paragraph(f"**Intercepto:** {intercept:.4f}")

    # Tabela de dados
    doc.add_page_break()
    add_data_table(doc, df)

    # Gráfico
    doc.add_heading("Gráfico 3D da Superfície", level=2)
    img = fig.to_image(format="png")
    doc.add_picture(BytesIO(img), width=Inches(6))

    buf = BytesIO()
    doc.save(buf)
    return buf

def generate_latex_doc(eq_latex, r2, r2_adj, rmse, mae,
                       mean_MR, std_MR, energy, degree,
                       intercept, df, fig):
    """
    Monta e retorna o conteúdo LaTeX (.tex) e os bytes da figura para zip.
    """
    lines = []
    # Preambulo
    lines.append(r"\documentclass{article}")
    lines.append(r"\usepackage[utf8]{inputenc}")
    lines.append(r"\usepackage{booktabs,graphicx}")
    lines.append(r"\begin{document}")
    lines.append(r"\section*{Relatório de Regressão}")

    # Configurações
    lines.append(r"\subsection*{Configurações}")
    lines.append(f"Tipo de energia: {energy}\\\\")
    if degree is not None:
        lines.append(f"Grau polinomial: {degree}\\\\")

    # Equação Ajustada
    lines.append(r"\subsection*{Equação Ajustada}")
    lines.append(eq_latex)

    # Indicadores Estatísticos
    lines.append(r"\subsection*{Indicadores Estatísticos}")
    lines.append(r"\begin{itemize}")
    lines.append(f"  \\item \\textbf{{R$^2$}}: {r2:.6f} (aprox. {r2*100:.2f}\\% explicado)")
    lines.append(f"  \\item \\textbf{{R$^2$ Ajustado}}: {r2_adj:.6f}")
    lines.append(f"  \\item \\textbf{{RMSE}}: {rmse:.4f} MPa")
    lines.append(f"  \\item \\textbf{{MAE}}: {mae:.4f} MPa")
    lines.append(f"  \\item \\textbf{{Média MR}}: {mean_MR:.4f} MPa")
    lines.append(f"  \\item \\textbf{{Desvio Padrão MR}}: {std_MR:.4f} MPa")
    lines.append(r"\end{itemize}")

    # Avaliação da Qualidade do Ajuste
    amp = df["MR"].max() - df["MR"].min()
    nrmse_range = rmse / amp if amp > 0 else float('nan')
    cv_rmse = rmse / mean_MR if mean_MR != 0 else float('nan')
    mae_pct = mae / mean_MR if mean_MR != 0 else float('nan')

    lines.append(r"\subsection*{Avaliação da Qualidade do Ajuste}")
    lines.append(r"\begin{itemize}")
    lines.append(f"  \\item \\textbf{{NRMSE_range}}: {nrmse_range:.2%}")
    lines.append(f"  \\item \\textbf{{CV(RMSE)}}: {cv_rmse:.2%}")
    lines.append(f"  \\item \\textbf{{MAE \\%}}: {mae_pct:.2%}")
    lines.append(r"\end{itemize}")

    # Intercepto
    lines.append(f"Intercepto: {intercept:.4f}\\\\")
    lines.append(r"\newpage")

    # Tabela de dados em LaTeX
    lines.append(r"\section*{Dados do Ensaio Triaxial}")
    cols = len(df.columns)
    lines.append(r"\begin{tabular}{" + "l" * cols + r"}")
    # Cabeçalho
    lines.append(" & ".join(df.columns) + " \\\\ \\midrule")
    # Linhas de dados
    for _, row in df.iterrows():
        vals = [str(v) for v in row.values]
        lines.append(" & ".join(vals) + " \\\\")
    lines.append(r"\end{tabular}")

    # Gráfico 3D
    lines.append(r"\section*{Gráfico 3D da Superfície}")
    lines.append(r"\includegraphics[width=\linewidth]{surface_plot.png}")
    lines.append(r"\end{document}")

    # Gera bytes da figura
    img_data = fig.to_image(format="png")
    tex_content = "\n".join(lines)
    return tex_content, img_data
