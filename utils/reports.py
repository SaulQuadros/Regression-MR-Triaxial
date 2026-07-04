import io
import math
import re
import numpy as np
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from mpl_toolkits.mplot3d import Axes3D
from matplotlib.gridspec import GridSpec
from docx import Document
from docx.shared import Inches, Pt
import zipfile

try:
    import streamlit as st
except ModuleNotFoundError:
    st = None

def get_export_view_offsets():
    if st is None:
        return 0, 0
    return st.session_state.get("azim_offset", 0), st.session_state.get("elev_offset", 0)

def export_plotly_figure_png(fig, azim_offset=0.0, elev_offset=0.0):
    if fig is None:
        return None
    # ----- surface trace -----
    surface_trace = None
    for tr in fig.data:
        if getattr(tr, "type", "") == "surface":
            surface_trace = tr
            break
    if surface_trace is None:
        return None

    Z_raw = np.array(surface_trace.z, dtype=float)
    X_raw = np.array(surface_trace.x, dtype=float)
    Y_raw = np.array(surface_trace.y, dtype=float)

    if X_raw.ndim == 1 and Y_raw.ndim == 1:
        nx = len(X_raw); ny = len(Y_raw)
        if Z_raw.shape == (ny, nx): Z = Z_raw
        elif Z_raw.shape == (nx, ny): Z = Z_raw.T
        else: Z = Z_raw.reshape(ny, nx)
        Xg, Yg = np.meshgrid(X_raw, Y_raw)
    else:
        Xg, Yg, Z = X_raw, Y_raw, Z_raw

    vmin, vmax = float(np.nanmin(Z)), float(np.nanmax(Z))
    fig_m = plt.figure(figsize=(7.0, 5.2), dpi=220)
    gs = GridSpec(1, 20, figure=fig_m)
    ax = fig_m.add_subplot(gs[0, :18], projection="3d")
    cax = fig_m.add_subplot(gs[0, 19])

    surf = ax.plot_surface(Xg, Yg, Z, cmap="viridis", linewidth=0, antialiased=True, shade=False, vmin=vmin, vmax=vmax)

    for tr in fig.data:
        if getattr(tr, "type", "") == "scatter3d":
            ax.scatter(np.array(tr.x), np.array(tr.y), np.array(tr.z), s=20, c="red")

    scene = getattr(fig.layout, "scene", {})
    ax.set_xlabel(r"$\sigma_3$ (MPa)"); ax.set_ylabel(r"$\sigma_d$ (MPa)"); ax.set_zlabel("MR (MPa)")

    try:
        cam = scene.camera
        eye = cam.eye
        ex, ey, ez = float(eye.x), float(eye.y), float(eye.z)
        az_plotly = math.degrees(math.atan2(ey, ex))
        az_mpl = (180.0 - az_plotly) + az_offset
        r = (ex**2 + ey**2 + ez**2) ** 0.5
        elev = math.degrees(math.asin(ez / r)) + elev_offset
        ax.view_init(elev=elev, azim=az_mpl)
    except:
        ax.view_init(elev=30 + elev_offset, azim=-60 + azim_offset)

    plt.colorbar(surf, cax=cax)
    out = io.BytesIO()
    fig_m.savefig(out, format="png", bbox_inches="tight")
    plt.close(fig_m)
    return out.getvalue()

def add_formatted_equation(doc, eq_text):
    eq = eq_text.strip().strip("$$")
    p = doc.add_paragraph(style="List Paragraph")
    i = 0
    while i < len(eq):
        ch = eq[i]
        if ch in ('^', '_'):
            is_sup = (ch == '^')
            i += 1
            if i < len(eq) and eq[i] == '{':
                i += 1
                content = ''
                while i < len(eq) and eq[i] != '}':
                    content += eq[i]
                    i += 1
                i += 1
            else:
                content = eq[i]
                i += 1
            run = p.add_run(content)
            if is_sup: run.font.superscript = True
            else: run.font.subscript = True
        elif ch == 'σ':
            run_sigma = p.add_run('σ')
            i += 1
            if i < len(eq) and eq[i] == '_':
                i += 1
                if i < len(eq) and eq[i] == '{':
                    i += 1
                    sub = ''
                    while i < len(eq) and eq[i] != '}':
                        sub += eq[i]
                        i += 1
                    i += 1
                else:
                    sub = eq[i]; i += 1
                run_sub = p.add_run(sub)
                run_sub.font.subscript = True
        else:
            p.add_run(ch); i += 1

def split_coefficient_label(label):
    """Separa um rótulo como 'k1' em base ('k') e subscrito ('1')."""
    match = re.match(r"^([^\d]+?)_?(\d*)$", label)
    if match:
        return match.group(1), match.group(2)
    return label, ""

def xml_escape(text):
    return str(text).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

DATA_TABLE_TITLE = "Anexo – Dados do Ensaio Triaxial"

# Cabeçalhos conhecidos renderizados em modo matemático no LaTeX
LATEX_COLUMN_MATH = {
    "σ3": r"$\sigma_3$",
    "σ₃": r"$\sigma_3$",
    "σd": r"$\sigma_d$",
    "σ_d": r"$\sigma_d$",
    "MR": r"$MR$",
    "θ": r"$\theta$",
}

def latex_column_header(name):
    return LATEX_COLUMN_MATH.get(str(name), latex_escape(str(name)))

def format_data_cell(value):
    if isinstance(value, (bool, np.bool_)):
        return str(value)
    if isinstance(value, (int, np.integer)):
        return str(int(value))
    if isinstance(value, (float, np.floating)):
        if np.isnan(value):
            return ""
        return f"{value:g}"
    return str(value)

def equation_to_pdf_markup(eq_text):
    """Converte a equação (σ_x, θ, ^{...}, _{...}, \\\\) em marcação inline do
    reportlab (<super>/<sub>/<br/>), preservando caracteres unicode."""
    eq = eq_text.strip().strip("$$")
    eq = eq.replace("\\\\", "")  # quebra de linha (usada no polinomial)
    eq = eq.replace("\\quad", "  ")
    out = []
    i = 0
    while i < len(eq):
        ch = eq[i]
        if ch in ("^", "_"):
            tag = "super" if ch == "^" else "sub"
            i += 1
            if i < len(eq) and eq[i] == "{":
                i += 1
                content = ""
                while i < len(eq) and eq[i] != "}":
                    content += eq[i]
                    i += 1
                i += 1  # pula '}'
            elif i < len(eq):
                content = eq[i]
                i += 1
            else:
                content = ""
            out.append(f"<{tag}>{xml_escape(content)}</{tag}>")
        elif ch == "":
            out.append("<br/>")
            i += 1
        else:
            out.append(xml_escape(ch))
            i += 1
    return "".join(out)

def coefficient_to_pdf_markup(label, value):
    base, sub = split_coefficient_label(label)
    symbol = f"{xml_escape(base)}<sub>{xml_escape(sub)}</sub>" if sub else f"<b>{xml_escape(base)}</b>"
    return f"{symbol} = {value:.6f}"

def add_coefficients_section(doc, coefficients):
    if not coefficients:
        return
    doc.add_heading("Coeficientes Calibrados", level=2)
    for label, value in coefficients:
        base, sub = split_coefficient_label(label)
        p = doc.add_paragraph(style="List Paragraph")
        p.add_run(base)
        if sub:
            sub_run = p.add_run(sub)
            sub_run.font.subscript = True
        p.add_run(f" = {value:.6f}")

def latex_coefficients_section(coefficients):
    if not coefficients:
        return []
    lines = [r"\subsection*{Coeficientes Calibrados}", r"\begin{itemize}"]
    for label, value in coefficients:
        base, sub = split_coefficient_label(label)
        symbol = rf"${latex_escape(base)}_{{{sub}}}$" if sub else rf"\textbf{{{latex_escape(base)}}}"
        lines.append(rf"\item {symbol}: ${value:.6f}$")
    lines.append(r"\end{itemize}")
    return lines

def add_key_value_section(doc, title, values):
    if not values:
        return
    doc.add_heading(title, level=2)
    for key, value in values.items():
        doc.add_paragraph(f"{key}: {value}", style="List Paragraph")

def add_data_table_section(doc, df, title=DATA_TABLE_TITLE):
    if df is None or getattr(df, "empty", True):
        return
    doc.add_page_break()
    doc.add_heading(title, level=2)
    columns = list(df.columns)
    table = doc.add_table(rows=1, cols=len(columns))
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    for j, col in enumerate(columns):
        cell = table.rows[0].cells[j]
        cell.text = str(col)
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for j, col in enumerate(columns):
            cells[j].text = format_data_cell(row[col])

def latex_data_table_section(df, title=DATA_TABLE_TITLE):
    if df is None or getattr(df, "empty", True):
        return []
    columns = list(df.columns)
    col_spec = "".join(["r"] * len(columns))
    header = " & ".join(latex_column_header(col) for col in columns) + r" \\"
    lines = [
        r"\clearpage",
        rf"\subsection*{{{latex_escape(title)}}}",
        rf"\begin{{longtable}}{{{col_spec}}}",
        r"\toprule", header, r"\midrule", r"\endhead",
    ]
    for _, row in df.iterrows():
        lines.append(" & ".join(latex_escape(format_data_cell(row[col])) for col in columns) + r" \\")
    lines.extend([r"\bottomrule", r"\end{longtable}"])
    return lines

def latex_escape(value):
    text = str(value)
    replacements = {
        "\\": r"\textbackslash{}",
        "&": r"\&",
        "%": r"\%",
        "$": r"\$",
        "#": r"\#",
        "_": r"\_",
        "{": r"\{",
        "}": r"\}",
        "~": r"\textasciitilde{}",
        "^": r"\textasciicircum{}",
    }
    return "".join(replacements.get(ch, ch) for ch in text)

def latex_key_value_section(title, values):
    if not values:
        return []
    lines = [rf"\subsection*{{{latex_escape(title)}}}", r"\begin{itemize}"]
    for key, value in values.items():
        lines.append(rf"\item \textbf{{{latex_escape(key)}}}: {latex_escape(value)}")
    lines.append(r"\end{itemize}")
    return lines

def quality_label(value, thresholds, labels):
    for threshold, label in zip(thresholds, labels):
        if value <= threshold:
            return label
    return labels[-1]

def statistical_metric_descriptions(metrics):
    return [
        (
            "R²",
            f"{metrics['r2']:.6f}",
            f"aproximadamente {metrics['r2'] * 100:.2f}% da variabilidade de MR é explicada pelo modelo.",
        ),
        (
            "R² ajustado",
            f"{metrics['r2_adj']:.6f}",
            "considera a quantidade de parâmetros usados no ajuste.",
        ),
        (
            "RMSE",
            f"{metrics['rmse']:.4f} MPa",
            "erro quadrático médio na unidade original do MR.",
        ),
        (
            "MAE",
            f"{metrics['mae']:.4f} MPa",
            "erro absoluto médio, menos sensível a erros extremos.",
        ),
        (
            "Média MR",
            f"{metrics['mean_y']:.4f} MPa",
            "valor médio observado no conjunto calibrado.",
        ),
        (
            "Desvio padrão MR",
            f"{metrics['std_y']:.4f} MPa",
            "dispersão dos valores de MR em torno da média.",
        ),
        (
            "Amplitude",
            f"{metrics['amplitude']:.4f} MPa",
            "diferença entre o maior e o menor MR observado.",
        ),
        (
            "MR máximo / mínimo",
            f"{metrics['max_y']:.4f} / {metrics['min_y']:.4f} MPa",
            "limites observados nos dados.",
        ),
    ]

def quality_metric_descriptions(metrics):
    labels_nrmse = ["Excelente (≤5%)", "Bom (≤10%)", "Insuficiente (>10%)"]
    labels_cv = ["Excelente (≤10%)", "Bom (≤20%)", "Insuficiente (>20%)"]
    return [
        (
            "NRMSE",
            f"{metrics['nrmse_range']:.2%}",
            quality_label(metrics["nrmse_range"], [0.05, 0.10], labels_nrmse),
        ),
        (
            "CV(RMSE)",
            f"{metrics['cv_rmse']:.2%}",
            quality_label(metrics["cv_rmse"], [0.10, 0.20], labels_cv),
        ),
        (
            "MAE %",
            f"{metrics['mae_pct']:.2%}",
            quality_label(metrics["mae_pct"], [0.10, 0.20], labels_cv),
        ),
    ]

def add_metric_description_section(doc, title, descriptions):
    doc.add_heading(title, level=2)
    for name, value, description in descriptions:
        doc.add_paragraph(f"{name}: {value} → {description}", style="List Paragraph")

def add_metric_sections_together(doc, sections):
    """Adiciona várias seções de métricas mantendo-as juntas na mesma página
    (evita que a última linha, ex.: MAE %, vaze para a página do gráfico)."""
    paragraphs = []
    for title, descriptions in sections:
        paragraphs.append(doc.add_heading(title, level=2))
        for name, value, description in descriptions:
            paragraphs.append(
                doc.add_paragraph(f"{name}: {value} → {description}", style="List Paragraph")
            )
    for paragraph in paragraphs:
        paragraph.paragraph_format.keep_together = True
    for paragraph in paragraphs[:-1]:
        paragraph.paragraph_format.keep_with_next = True

def latex_metric_description_section(title, descriptions):
    lines = [rf"\subsection*{{{latex_escape(title)}}}", r"\begin{itemize}"]
    for name, value, description in descriptions:
        lines.append(
            rf"\item \textbf{{{latex_escape(name)}}}: {latex_escape(value)} $\rightarrow$ {latex_escape(description)}"
        )
    lines.append(r"\end{itemize}")
    return lines

def generate_word_doc(model, metrics, df, fig, energy, traceability=None, modeling_metadata=None):
    doc = Document()
    for style in ['Normal', 'List Paragraph', 'Heading 1', 'Heading 2']:
        if style in doc.styles: doc.styles[style].font.size = Pt(12)

    doc.add_heading("Relatório de Regressão", level=1)
    add_key_value_section(doc, "Rastreabilidade", traceability)

    if modeling_metadata is None:
        modeling_metadata = {"Modelo ajustado": model.name, "Energia": energy}
    add_key_value_section(doc, "Configurações de Modelagem", modeling_metadata)

    doc.add_heading("Equação Ajustada", level=2)
    add_formatted_equation(doc, model.get_equation())
    equation_note = model.get_equation_note()
    if equation_note:
        add_formatted_equation(doc, equation_note)
    add_coefficients_section(doc, model.get_coefficients())

    statistical_descriptions = statistical_metric_descriptions(metrics)
    statistical_descriptions.append(
        ("Intercepto", f"{model.intercept:.4f} MPa", "termo constante do modelo ajustado.")
    )
    add_metric_sections_together(doc, [
        ("Indicadores Estatísticos", statistical_descriptions),
        ("Qualidade do Ajuste", quality_metric_descriptions(metrics)),
    ])

    doc.add_page_break()
    doc.add_heading("Gráfico 3D", level=2)
    azim_offset, elev_offset = get_export_view_offsets()
    img_data = export_plotly_figure_png(fig, azim_offset, elev_offset)
    if img_data: doc.add_picture(io.BytesIO(img_data), width=Inches(6))

    add_data_table_section(doc, df)

    buf = io.BytesIO()
    doc.save(buf)
    return buf

def generate_latex_zip(model, metrics, df, fig, energy, traceability=None, modeling_metadata=None):
    if modeling_metadata is None:
        modeling_metadata = {"Modelo ajustado": model.name, "Energia": energy}

    lines = [
        r"\documentclass{article}", r"\usepackage[utf8]{inputenc}", r"\usepackage{booktabs,graphicx,longtable}",
        r"\begin{document}", r"\section*{Relatório de Regressão}",
    ]
    lines.extend(latex_key_value_section("Rastreabilidade", traceability))
    lines.extend(latex_key_value_section("Configurações de Modelagem", modeling_metadata))
    lines.extend([
        r"\subsection*{Equação}", model.get_equation(),
    ])
    equation_note = model.get_equation_note()
    if equation_note:
        lines.append(f"$${equation_note}$$")
    lines.extend(latex_coefficients_section(model.get_coefficients()))
    statistical_descriptions = statistical_metric_descriptions(metrics)
    statistical_descriptions.append(
        ("Intercepto", f"{model.intercept:.4f} MPa", "termo constante do modelo ajustado.")
    )
    lines.extend(latex_metric_description_section("Indicadores Estatísticos", statistical_descriptions))
    lines.extend(latex_metric_description_section("Qualidade do Ajuste", quality_metric_descriptions(metrics)))
    lines.extend([
        r"\clearpage",
        r"\subsection*{Gráfico 3D}",
        r"\begin{figure}[htbp]",
        r"\centering",
        r"\includegraphics[width=0.95\textwidth]{surface_plot.png}",
        r"\caption{Superfície ajustada de MR em função de $\sigma_3$ e $\sigma_d$.}",
        r"\end{figure}",
    ])
    lines.extend(latex_data_table_section(df))
    lines.append(r"\end{document}")
    tex_content = "\n".join(lines)
    azim_offset, elev_offset = get_export_view_offsets()
    img_data = export_plotly_figure_png(fig, azim_offset, elev_offset)
    
    zip_buf = io.BytesIO()
    with zipfile.ZipFile(zip_buf, "w") as zf:
        zf.writestr("main.tex", tex_content)
        if img_data: zf.writestr("surface_plot.png", img_data)
    zip_buf.seek(0)
    return zip_buf, tex_content

def _register_pdf_fonts():
    """Registra a fonte DejaVu Sans (distribuída com o matplotlib) para o
    reportlab, garantindo suporte a σ, θ, subscritos e acentos no PDF."""
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    import matplotlib.font_manager as fm

    if "DejaVuSans" in pdfmetrics.getRegisteredFontNames():
        return
    pdfmetrics.registerFont(TTFont("DejaVuSans", fm.findfont("DejaVu Sans")))
    try:
        bold_path = fm.findfont("DejaVu Sans:bold")
    except Exception:
        bold_path = fm.findfont("DejaVu Sans")
    pdfmetrics.registerFont(TTFont("DejaVuSans-Bold", bold_path))
    pdfmetrics.registerFontFamily("DejaVuSans", normal="DejaVuSans", bold="DejaVuSans-Bold")

def generate_pdf_doc(model, metrics, df, fig, energy, traceability=None, modeling_metadata=None):
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.utils import ImageReader
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Image as RLImage, Table, TableStyle, PageBreak, KeepTogether,
    )

    _register_pdf_fonts()
    if modeling_metadata is None:
        modeling_metadata = {"Modelo ajustado": model.name, "Energia": energy}

    body = ParagraphStyle("body", fontName="DejaVuSans", fontSize=10, leading=15)
    h1 = ParagraphStyle("h1", fontName="DejaVuSans-Bold", fontSize=16, leading=20, spaceAfter=10)
    h2 = ParagraphStyle("h2", fontName="DejaVuSans-Bold", fontSize=12, leading=16,
                        spaceBefore=12, spaceAfter=6)
    eq_style = ParagraphStyle("eq", fontName="DejaVuSans", fontSize=12, leading=18,
                              spaceBefore=4, spaceAfter=8)

    story = [Paragraph("Relatório de Regressão", h1)]

    def kv_section(title, values):
        if not values:
            return
        story.append(Paragraph(xml_escape(title), h2))
        for key, value in values.items():
            story.append(Paragraph(f"<b>{xml_escape(key)}</b>: {xml_escape(value)}", body))

    def metric_flowables(title, descriptions):
        flowables = [Paragraph(xml_escape(title), h2)]
        for name, value, description in descriptions:
            flowables.append(Paragraph(
                f"<b>{xml_escape(name)}</b>: {xml_escape(value)} → {xml_escape(description)}", body))
        return flowables

    kv_section("Rastreabilidade", traceability)
    kv_section("Configurações de Modelagem", modeling_metadata)

    story.append(Paragraph("Equação Ajustada", h2))
    story.append(Paragraph(equation_to_pdf_markup(model.get_equation()), eq_style))
    equation_note = model.get_equation_note()
    if equation_note:
        story.append(Paragraph(equation_to_pdf_markup(equation_note), eq_style))

    coefficients = model.get_coefficients()
    if coefficients:
        story.append(Paragraph("Coeficientes Calibrados", h2))
        for label, value in coefficients:
            story.append(Paragraph(coefficient_to_pdf_markup(label, value), body))

    statistical_descriptions = statistical_metric_descriptions(metrics)
    statistical_descriptions.append(
        ("Intercepto", f"{model.intercept:.4f} MPa", "termo constante do modelo ajustado.")
    )
    # Mantém as duas seções de métricas juntas na mesma página
    story.append(KeepTogether(
        metric_flowables("Indicadores Estatísticos", statistical_descriptions)
        + metric_flowables("Qualidade do Ajuste", quality_metric_descriptions(metrics))
    ))

    # Gráfico sempre em nova página, para não dividir página com as métricas
    azim_offset, elev_offset = get_export_view_offsets()
    img_data = export_plotly_figure_png(fig, azim_offset, elev_offset)
    if img_data:
        iw, ih = ImageReader(io.BytesIO(img_data)).getSize()
        width = 16 * cm
        story.append(PageBreak())
        story.append(Paragraph("Gráfico 3D", h2))
        story.append(RLImage(io.BytesIO(img_data), width=width, height=width * ih / iw))

    if df is not None and not getattr(df, "empty", True):
        story.append(PageBreak())
        story.append(Paragraph(xml_escape(DATA_TABLE_TITLE), h2))
        columns = list(df.columns)
        table_data = [[xml_escape(col) for col in columns]]
        for _, row in df.iterrows():
            table_data.append([format_data_cell(row[col]) for col in columns])
        data_table = Table(table_data, repeatRows=1)
        data_table.setStyle(TableStyle([
            ("FONTNAME", (0, 0), (-1, -1), "DejaVuSans"),
            ("FONTNAME", (0, 0), (-1, 0), "DejaVuSans-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), 9),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
            ("ALIGN", (0, 0), (-1, -1), "CENTER"),
            ("TOPPADDING", (0, 0), (-1, -1), 3),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
        ]))
        story.append(data_table)

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        leftMargin=2 * cm, rightMargin=2 * cm, topMargin=2 * cm, bottomMargin=2 * cm,
    )
    doc.build(story)
    buf.seek(0)
    return buf
