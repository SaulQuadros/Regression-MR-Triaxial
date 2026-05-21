import io
import math
import numpy as np
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from mpl_toolkits.mplot3d import Axes3D
from matplotlib.gridspec import GridSpec
from docx import Document
from docx.shared import Inches, Pt
import streamlit as st
import zipfile

def export_plotly_figure_png(fig, azim_offset=0.0, elev_offset=0.0):
    # ----- surface trace -----
    surface_trace = None
    for tr in fig.data:
        if getattr(tr, "type", "") == "surface":
            surface_trace = tr
            break
    if surface_trace is None:
        return None

    Z_raw = np.array(surface_trace.z, dtype=float)
    X_raw = np.array(surface_trace.x, dtype=float)
    Y_raw = np.array(surface_trace.y, dtype=float)

    if X_raw.ndim == 1 and Y_raw.ndim == 1:
        nx = len(X_raw); ny = len(Y_raw)
        if Z_raw.shape == (ny, nx): Z = Z_raw
        elif Z_raw.shape == (nx, ny): Z = Z_raw.T
        else: Z = Z_raw.reshape(ny, nx)
        Xg, Yg = np.meshgrid(X_raw, Y_raw)
    else:
        Xg, Yg, Z = X_raw, Y_raw, Z_raw

    vmin, vmax = float(np.nanmin(Z)), float(np.nanmax(Z))
    fig_m = plt.figure(figsize=(7.0, 5.2), dpi=220)
    gs = GridSpec(1, 20, figure=fig_m)
    ax = fig_m.add_subplot(gs[0, :18], projection="3d")
    cax = fig_m.add_subplot(gs[0, 19])

    surf = ax.plot_surface(Xg, Yg, Z, cmap="viridis", linewidth=0, antialiased=True, shade=False, vmin=vmin, vmax=vmax)

    for tr in fig.data:
        if getattr(tr, "type", "") == "scatter3d":
            ax.scatter(np.array(tr.x), np.array(tr.y), np.array(tr.z), s=20, c="red")

    scene = getattr(fig.layout, "scene", {})
    ax.set_xlabel("σ₃ (MPa)"); ax.set_ylabel("σ_d (MPa)"); ax.set_zlabel("MR (MPa)")

    try:
        cam = scene.camera
        eye = cam.eye
        ex, ey, ez = float(eye.x), float(eye.y), float(eye.z)
        az_plotly = math.degrees(math.atan2(ey, ex))
        az_mpl = (180.0 - az_plotly) + az_offset
        r = (ex**2 + ey**2 + ez**2) ** 0.5
        elev = math.degrees(math.asin(ez / r)) + elev_offset
        ax.view_init(elev=elev, azim=az_mpl)
    except:
        ax.view_init(elev=30, azim=-60)

    plt.colorbar(surf, cax=cax)
    out = io.BytesIO()
    fig_m.savefig(out, format="png", bbox_inches="tight")
    plt.close(fig_m)
    return out.getvalue()

def add_formatted_equation(doc, eq_text):
    eq = eq_text.strip().strip("$$")
    p = doc.add_paragraph(style="List Paragraph")
    i = 0
    while i < len(eq):
        ch = eq[i]
        if ch in ('^', '_'):
            is_sup = (ch == '^')
            i += 1
            if i < len(eq) and eq[i] == '{':
                i += 1
                content = ''
                while i < len(eq) and eq[i] != '}':
                    content += eq[i]
                    i += 1
                i += 1
            else:
                content = eq[i]
                i += 1
            run = p.add_run(content)
            if is_sup: run.font.superscript = True
            else: run.font.subscript = True
        elif ch == 'σ':
            run_sigma = p.add_run('σ')
            i += 1
            if i < len(eq) and eq[i] == '_':
                i += 1
                if i < len(eq) and eq[i] == '{':
                    i += 1
                    sub = ''
                    while i < len(eq) and eq[i] != '}':
                        sub += eq[i]
                        i += 1
                    i += 1
                else:
                    sub = eq[i]; i += 1
                run_sub = p.add_run(sub)
                run_sub.font.subscript = True
        else:
            p.add_run(ch); i += 1

def generate_word_doc(model, metrics, df, fig, energy):
    doc = Document()
    for style in ['Normal', 'List Paragraph', 'Heading 1', 'Heading 2']:
        if style in doc.styles: doc.styles[style].font.size = Pt(12)

    doc.add_heading("Relatório de Regressão", level=1)
    doc.add_heading("Configurações", level=2)
    doc.add_paragraph(f"Modelo: {model.name}", style="List Paragraph")
    doc.add_paragraph(f"Energia: {energy}", style="List Paragraph")

    doc.add_heading("Equação Ajustada", level=2)
    add_formatted_equation(doc, model.get_equation())

    doc.add_heading("Indicadores Estatísticos", level=2)
    doc.add_paragraph(f"R² = {metrics['r2']:.6f}", style="List Paragraph")
    doc.add_paragraph(f"R² Ajustado = {metrics['r2_adj']:.6f}", style="List Paragraph")
    doc.add_paragraph(f"RMSE = {metrics['rmse']:.4f} MPa", style="List Paragraph")
    doc.add_paragraph(f"MAE = {metrics['mae']:.4f} MPa", style="List Paragraph")
    doc.add_paragraph(f"Intercepto = {model.intercept:.4f} MPa", style="List Paragraph")

    doc.add_heading("Qualidade do Ajuste", level=2)
    doc.add_paragraph(f"NRMSE = {metrics['nrmse_range']:.2%}", style="List Paragraph")
    doc.add_paragraph(f"CV(RMSE) = {metrics['cv_rmse']:.2%}", style="List Paragraph")

    doc.add_page_break()
    doc.add_heading("Gráfico 3D", level=2)
    img_data = export_plotly_figure_png(fig, st.session_state.get('azim_offset', 0), st.session_state.get('elev_offset', 0))
    if img_data: doc.add_picture(io.BytesIO(img_data), width=Inches(6))

    buf = io.BytesIO()
    doc.save(buf)
    return buf

def generate_latex_zip(model, metrics, df, fig, energy):
    lines = [
        r"\documentclass{article}", r"\usepackage[utf8]{inputenc}", r"\usepackage{booktabs,graphicx}",
        r"\begin{document}", r"\section*{Relatório de Regressão}",
        f"Modelo: {model.name}\\", f"Energia: {energy}\\",
        r"\subsection*{Equação}", model.get_equation(),
        r"\subsection*{Métricas}", r"\begin{itemize}",
        f"\\item R$^2$: {metrics['r2']:.6f}", f"\\item RMSE: {metrics['rmse']:.4f} MPa",
        r"\end{itemize}", r"\end{document}"
    ]
    tex_content = "\n".join(lines)
    img_data = export_plotly_figure_png(fig, st.session_state.get('azim_offset', 0), st.session_state.get('elev_offset', 0))
    
    zip_buf = io.BytesIO()
    with zipfile.ZipFile(zip_buf, "w") as zf:
        zf.writestr("main.tex", tex_content)
        if img_data: zf.writestr("surface_plot.png", img_data)
    zip_buf.seek(0)
    return zip_buf, tex_content
