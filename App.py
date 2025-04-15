#!/usr/bin/env python
# coding: utf-8

# In[ ]:


import streamlit as st
import pandas as pd
import numpy as np
from io import BytesIO
from sklearn.preprocessing import PolynomialFeatures
from sklearn.linear_model import LinearRegression
from sklearn.metrics import r2_score, mean_squared_error, mean_absolute_error
import plotly.graph_objs as go
from docx import Document
from docx.shared import Inches
import base64

# Função para calcular R² ajustado
def adjusted_r2(r2, n, p):
    return 1 - ((1 - r2) * (n - 1)) / (n - p - 1)

# Função para construir string da equação em LaTeX
def build_latex_equation(coefs, intercept, feature_names):
    # Inicia com o intercepto
    equation = f"{intercept:.4f}"
    # Cada coeficiente e respectivo termo
    for coef, term in zip(coefs[1:], feature_names[1:]):
        # Ajusta o sinal
        sign = " + " if coef >= 0 else " - "
        equation += sign + f"{abs(coef):.4f} " + term.replace(" ", "")
    return "$$ MR = " + equation + " $$"

# Função para criar o gráfico 3D usando Plotly
def plot_3d_surface(df, model, poly, energy_col):
    # Gera uma grade para plotagem (usando os valores de σ3 e σd)
    sigma3_range = np.linspace(df["σ3"].min(), df["σ3"].max(), 30)
    sigmad_range = np.linspace(df["σd"].min(), df["σd"].max(), 30)
    sigma3_grid, sigmad_grid = np.meshgrid(sigma3_range, sigmad_range)
    X_grid = np.c_[sigma3_grid.ravel(), sigmad_grid.ravel()]
    X_poly_grid = poly.transform(X_grid)
    MR_pred = model.predict(X_poly_grid).reshape(sigma3_grid.shape)

    fig = go.Figure(data=[go.Surface(x=sigma3_grid, y=sigmad_grid, z=MR_pred, colorscale='Viridis')])
    # Adiciona pontos dos dados
    fig.add_trace(go.Scatter3d(
        x=df["σ3"],
        y=df["σd"],
        z=df[energy_col],
        mode='markers',
        marker=dict(size=5, color='red'),
        name="Dados"
    ))
    fig.update_layout(
        scene = dict(
            xaxis_title='σ₃ (MPa)',
            yaxis_title='σ_d (MPa)',
            zaxis_title='MR (MPa)'
        ),
        margin=dict(l=0, r=0, b=0, t=30)
    )
    return fig

# Função para gerar interpretação dos indicadores
def interpret_metrics(r2, r2_adj, rmse, mae):
    interpretation = f"**R²:** {r2:.6f}. Este valor indica que aproximadamente {r2*100:.2f}% da variabilidade dos dados de MR é explicada pelo modelo.\n\n"
    interpretation += f"**R² Ajustado:** {r2_adj:.6f}. Essa métrica penaliza o uso excessivo de termos. A alta similaridade com o R² mostra que o modelo não sofreu superajuste significativo.\n\n"
    interpretation += f"**RMSE:** {rmse:.4f} MPa. Esse valor indica que, em média, a previsão difere dos valores observados em cerca de {rmse:.4f} MPa (sensível a erros grandes).\n\n"
    interpretation += f"**MAE:** {mae:.4f} MPa. Essa métrica indica que, em média, o erro absoluto entre o valor previsto e o real é de {mae:.4f} MPa.\n\n"
    return interpretation

# Função para gerar documento Word com os resultados
def generate_word_doc(equation_latex, metrics_text, fig, energy_type, degree):
    document = Document()
    document.add_heading("Relatório de Regressão Polinomial", level=1)
    document.add_heading("Configurações do Modelo", level=2)
    document.add_paragraph(f"Tipo de energia: {energy_type}")
    document.add_paragraph(f"Grau da equação polinomial: {degree}")
    
    document.add_heading("Equação de Regressão", level=2)
    document.add_paragraph("A equação ajustada é apresentada abaixo (em LaTeX):")
    document.add_paragraph(equation_latex)
    
    document.add_heading("Indicadores Estatísticos", level=2)
    document.add_paragraph(metrics_text)
    
    document.add_heading("Gráfico 3D da Superfície", level=2)
    # Salva o gráfico como imagem e adiciona ao Word
    img_bytes = fig.to_image(format="png")
    image_stream = BytesIO(img_bytes)
    document.add_picture(image_stream, width=Inches(6))
    
    # Salva o documento em um buffer
    doc_buffer = BytesIO()
    document.save(doc_buffer)
    return doc_buffer

# --- Aplicativo Streamlit ---
st.set_page_config(page_title="Regressão Polinomial para MR", layout="wide")
st.title("Repressão Polinomial para Módulo de Resiliência (MR)")
st.markdown("Este app permite o upload de uma tabela com os parâmetros *σ₃* (tensão confinante), *σ_d* (tensão desvio) e *MR* (módulo de resiliência) e ajusta uma equação polinomial (grau 2 a 6) por regressão não linear.")

# Upload da tabela
uploaded_file = st.file_uploader("Faça o upload da tabela (CSV ou Excel)", type=["csv", "xlsx"])
if uploaded_file is not None:
    # Tenta ler o arquivo
    try:
        if uploaded_file.name.endswith('.csv'):
            df = pd.read_csv(uploaded_file, decimal=",")
        else:
            df = pd.read_excel(uploaded_file)
        st.write("### Dados Carregados")
        st.dataframe(df)
    except Exception as e:
        st.error("Erro ao carregar o arquivo!")
else:
    st.info("Por favor, faça o upload de um arquivo para continuar.")

# Apenas procede se os dados estiverem carregados
if uploaded_file is not None:
    # Opções do usuário
    st.sidebar.header("Configurações do Modelo")
    degree = st.sidebar.selectbox("Selecione o grau da equação polinomial", options=[2,3,4,5,6], index=0)
    energy_type = st.sidebar.selectbox("Selecione o tipo de energia", options=["Normal", "Intermediária", "Modificada"], index=0)
    
    # Define a coluna a ser usada conforme o tipo selecionado:
    # Supondo que para "Normal" a coluna seja "MR".
    energy_col = "MR"  # Se a tabela tiver colunas "MR", "MR_I", "MR_M", adapte conforme necessário.
    
    if st.button("Calcular"):
        # Preparação dos dados: considera as colunas "σ3", "σd" e a coluna definida em energy_col.
        try:
            df = df.rename(columns=lambda x: x.strip())
            X_data = df[["σ3", "σd"]].values
            y_data = df[energy_col].values
        except Exception as e:
            st.error("Certifique-se de que a tabela possua as colunas 'σ3', 'σd' e 'MR' (ou equivalente).")
        
        # Cria os termos polinomiais
        poly = PolynomialFeatures(degree=degree)
        X_poly = poly.fit_transform(X_data)
        
        # Ajuste da regressão linear (mínimos quadrados)
        model = LinearRegression()
        model.fit(X_poly, y_data)
        y_pred = model.predict(X_poly)
        
        # Cálculo dos indicadores
        r2 = r2_score(y_data, y_pred)
        n_obs = len(y_data)
        p_pred = X_poly.shape[1] - 1  # número de preditores (excluindo o intercepto)
        r2_adj = adjusted_r2(r2, n_obs, p_pred)
        rmse = np.sqrt(mean_squared_error(y_data, y_pred))
        mae = mean_absolute_error(y_data, y_pred)
        
        # Constrói a equação em LaTeX
        feature_names = poly.get_feature_names_out(["σ₃", "σ_d"])
        equation_latex = build_latex_equation(model.coef_, model.intercept_, feature_names)
        
        # Gera interpretações dos indicadores (exceto intercepto)
        metrics_interpretation = interpret_metrics(r2, r2_adj, rmse, mae)
        
        st.write("### Equação de Regressão (LaTeX)")
        st.latex(equation_latex.strip("$$"))
        
        st.write("### Indicadores Estatísticos")
        st.markdown(metrics_interpretation)
        st.write(f"**Intercepto:** {model.intercept_:.4f}")
        
        # Exibe o gráfico 3D
        st.write("### Gráfico 3D da Superfície")
        fig = plot_3d_surface(df, model, poly, energy_col)
        st.plotly_chart(fig, use_container_width=True)
        
        # Botão de download do Word
        if st.button("Salvar Word"):
            doc_buffer = generate_word_doc(equation_latex, metrics_interpretation, fig, energy_type, degree)
            doc_buffer.seek(0)
            b64 = base64.b64encode(doc_buffer.read()).decode()
            href = f'<a href="data:application/octet-stream;base64,{b64}" download="Relatorio_Regressao.docx">Clique aqui para baixar o arquivo Word</a>'
            st.markdown(href, unsafe_allow_html=True)

